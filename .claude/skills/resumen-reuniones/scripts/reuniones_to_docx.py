#!/usr/bin/env python3
"""
Genera un documento Word con resumen de reuniones a partir de datos JSON.

Uso:
    python3 scripts/reuniones_to_docx.py <json_input> <docx_output>

Input JSON:
    {
        "entidad": "IDT",
        "periodo_inicio": "2026-02-01",
        "periodo_fin": "2026-02-28",
        "total_reuniones": 8,
        "con_transcripcion": 5,
        "solo_calendario": 3,
        "reuniones": [
            {
                "nombre": "Sprint Planning",
                "fecha": "2026-02-05",
                "hora": "10:00",
                "asistentes": ["persona1@ent.gov.co", "persona2@ent.gov.co"],
                "resumen": "Se planificaron las tareas del sprint...",
                "puntos_clave": ["Punto 1", "Punto 2"],
                "acuerdos": ["Acuerdo 1", "Acuerdo 2"],
                "tiene_transcripcion": true
            }
        ],
        "tabla_resumen": [
            {
                "num": 1,
                "fecha": "2026-02-05",
                "nombre": "Sprint Planning",
                "participantes": "5 personas",
                "temas": "Planificacion sprint 10"
            }
        ]
    }
"""

import json
import os
import sys
from datetime import datetime

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Inches, Pt, RGBColor


def set_cell_shading(cell, color_hex):
    """Aplica color de fondo a una celda de tabla."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    shading = OxmlElement("w:shd")
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:color"), "auto")
    shading.set(qn("w:fill"), color_hex)
    cell._tc.get_or_add_tcPr().append(shading)


def generate_docx(data, output_path):
    """Genera el documento Word a partir de los datos JSON."""
    entidad = data.get("entidad", "")
    periodo_inicio = data.get("periodo_inicio", "")
    periodo_fin = data.get("periodo_fin", "")
    total_reuniones = data.get("total_reuniones", 0)
    con_transcripcion = data.get("con_transcripcion", 0)
    solo_calendario = data.get("solo_calendario", 0)
    reuniones = data.get("reuniones", [])
    tabla_resumen = data.get("tabla_resumen", [])

    doc = Document()

    # Estilos del documento
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)

    # Titulo
    title = doc.add_heading(level=0)
    run = title.add_run(f"RESUMEN DE REUNIONES - {entidad}")
    run.font.color.rgb = RGBColor(0, 51, 102)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Subtitulo
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(f"Periodo: {periodo_inicio} a {periodo_fin}")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(100, 100, 100)

    # Estadisticas
    stats_para = doc.add_paragraph()
    stats_para.add_run(f"Total de reuniones: {total_reuniones}").bold = True
    stats_para.add_run(f"  |  Con transcripcion: {con_transcripcion}")
    stats_para.add_run(f"  |  Solo calendario: {solo_calendario}")

    doc.add_paragraph()

    # Tabla resumen
    if tabla_resumen:
        doc.add_heading("Tabla Resumen", level=1)
        headers = ["#", "Fecha", "Reunion", "Participantes", "Temas Principales"]
        table = doc.add_table(rows=1, cols=len(headers))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = "Table Grid"

        # Header
        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = header
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(255, 255, 255)
                    run.font.size = Pt(9)
            set_cell_shading(cell, "003366")

        # Filas
        for row_data in tabla_resumen:
            row = table.add_row()
            values = [
                str(row_data.get("num", "")),
                str(row_data.get("fecha", "")),
                str(row_data.get("nombre", "")),
                str(row_data.get("participantes", "")),
                str(row_data.get("temas", "")),
            ]
            for i, val in enumerate(values):
                cell = row.cells[i]
                cell.text = val
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(9)

        # Ajustar anchos
        widths = [Cm(1.5), Cm(2.5), Cm(5), Cm(3), Cm(6)]
        for row in table.rows:
            for i, width in enumerate(widths):
                row.cells[i].width = width

        doc.add_paragraph()

    # Detalle por reunion
    reuniones_con_trans = [r for r in reuniones if r.get("tiene_transcripcion")]
    reuniones_sin_trans = [r for r in reuniones if not r.get("tiene_transcripcion")]

    if reuniones_con_trans:
        doc.add_heading("Detalle de Reuniones", level=1)

        for reunion in reuniones_con_trans:
            nombre = reunion.get("nombre", "Sin titulo")
            fecha = reunion.get("fecha", "")
            hora = reunion.get("hora", "")
            asistentes = reunion.get("asistentes", [])
            resumen = reunion.get("resumen", "")
            puntos_clave = reunion.get("puntos_clave", [])
            acuerdos = reunion.get("acuerdos", [])

            # Heading de reunion
            doc.add_heading(nombre, level=2)

            # Metadata
            meta = doc.add_paragraph()
            meta.add_run("Fecha y hora: ").bold = True
            meta.add_run(f"{fecha} {hora}")
            meta.add_run("\n")
            meta.add_run("Asistentes: ").bold = True
            if isinstance(asistentes, list):
                meta.add_run(", ".join(asistentes))
            else:
                meta.add_run(str(asistentes))

            # Resumen
            if resumen:
                res_heading = doc.add_paragraph()
                res_heading.add_run("Resumen de la sesion:").bold = True
                doc.add_paragraph(resumen)

            # Puntos clave
            if puntos_clave:
                pk_heading = doc.add_paragraph()
                pk_heading.add_run("Puntos clave:").bold = True
                for punto in puntos_clave:
                    doc.add_paragraph(punto, style="List Bullet")

            # Acuerdos
            if acuerdos:
                ac_heading = doc.add_paragraph()
                ac_heading.add_run("Acuerdos y compromisos:").bold = True
                for acuerdo in acuerdos:
                    doc.add_paragraph(acuerdo, style="List Bullet")

    # Reuniones sin transcripcion
    if reuniones_sin_trans:
        doc.add_heading("Reuniones sin transcripcion (datos de calendario)", level=1)
        doc.add_paragraph(
            "Las siguientes reuniones fueron identificadas en el calendario "
            "institucional pero no cuentan con transcripcion disponible:"
        )

        headers = ["#", "Fecha", "Hora", "Reunion", "Participantes"]
        table = doc.add_table(rows=1, cols=len(headers))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = "Table Grid"

        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = header
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(255, 255, 255)
                    run.font.size = Pt(9)
            set_cell_shading(cell, "666666")

        for idx, reunion in enumerate(reuniones_sin_trans, 1):
            row = table.add_row()
            asistentes = reunion.get("asistentes", [])
            if isinstance(asistentes, list):
                asistentes_str = f"{len(asistentes)} personas"
            else:
                asistentes_str = str(asistentes)
            values = [
                str(idx),
                str(reunion.get("fecha", "")),
                str(reunion.get("hora", "")),
                str(reunion.get("nombre", "")),
                asistentes_str,
            ]
            for i, val in enumerate(values):
                cell = row.cells[i]
                cell.text = val
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(9)

    # Footer con timestamp
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run(f"Generado automaticamente el {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(150, 150, 150)
    run.italic = True

    os.makedirs(os.path.dirname(output_path) or ".", exist_ok=True)
    doc.save(output_path)
    print(f"DOCX generado: {output_path} ({total_reuniones} reuniones)", file=sys.stderr)


def main():
    if len(sys.argv) < 3:
        print("Uso: python3 scripts/reuniones_to_docx.py <json_input> <docx_output>")
        sys.exit(1)

    json_input = sys.argv[1]
    docx_output = sys.argv[2]

    with open(json_input, "r", encoding="utf-8") as f:
        data = json.load(f)

    generate_docx(data, docx_output)


if __name__ == "__main__":
    main()
