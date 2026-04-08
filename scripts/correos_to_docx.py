#!/usr/bin/env python3
"""
Genera un documento Word con correos enviados, mostrando el contenido completo
de cada correo con formato similar a una impresion de correo.

Uso:
    python3 scripts/correos_to_docx.py <json_input> <docx_output>

Input JSON:
    {
        "entidad": "IDT",
        "periodo_inicio": "2026-02-01",
        "periodo_fin": "2026-02-28",
        "cuenta": "usuario@entidad.gov.co",
        "total": 15,
        "correos": [
            {
                "num": 1,
                "fecha": "2026-02-05",
                "asunto": "RE: Error modulo X",
                "de": "usuario@entidad.gov.co",
                "para": "destinatario@entidad.gov.co",
                "cc": "otro@entidad.gov.co",
                "cuerpo": "Texto completo del correo...",
                "hilo": [
                    {
                        "de": "destinatario@entidad.gov.co",
                        "para": "usuario@entidad.gov.co",
                        "fecha": "2026-02-04",
                        "cuerpo": "Mensaje anterior del hilo..."
                    }
                ]
            }
        ]
    }
"""

import json
import os
import sys
from datetime import datetime

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


def set_document_language(doc, lang_code="es-CO"):
    """Configura el idioma del documento Word a español (Colombia)."""
    styles_element = doc.styles.element
    rpr_default = styles_element.find(qn("w:docDefaults"))
    if rpr_default is None:
        rpr_default = OxmlElement("w:docDefaults")
        styles_element.insert(0, rpr_default)
    rpr = rpr_default.find(qn("w:rPrDefault"))
    if rpr is None:
        rpr = OxmlElement("w:rPrDefault")
        rpr_default.append(rpr)
    rpr_inner = rpr.find(qn("w:rPr"))
    if rpr_inner is None:
        rpr_inner = OxmlElement("w:rPr")
        rpr.append(rpr_inner)
    lang = rpr_inner.find(qn("w:lang"))
    if lang is None:
        lang = OxmlElement("w:lang")
        rpr_inner.append(lang)
    lang.set(qn("w:val"), lang_code)
    lang.set(qn("w:eastAsia"), lang_code)
    lang.set(qn("w:bidi"), lang_code)


def add_horizontal_line(doc):
    """Agrega una línea horizontal al documento."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    pPr = p._p.get_or_add_pPr()
    pBdr = pPr.makeelement(qn("w:pBdr"), {})
    bottom = pBdr.makeelement(
        qn("w:bottom"),
        {
            qn("w:val"): "single",
            qn("w:sz"): "6",
            qn("w:space"): "1",
            qn("w:color"): "999999",
        },
    )
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_metadata_line(doc, label, value, indent=False):
    """Agrega una línea de metadatos (De:, Para:, etc.)."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    if indent:
        p.paragraph_format.left_indent = Inches(0.3)
    run_label = p.add_run(f"{label}: ")
    run_label.bold = True
    run_label.font.size = Pt(10)
    run_label.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    run_value = p.add_run(str(value) if value else "")
    run_value.font.size = Pt(10)
    run_value.font.color.rgb = RGBColor(0x55, 0x55, 0x55)


def add_body_text(doc, body, indent=False):
    """Agrega el cuerpo del correo como texto."""
    if not body:
        return
    lines = str(body).split("\n")
    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        if indent:
            p.paragraph_format.left_indent = Inches(0.3)
        run = p.add_run(line)
        run.font.size = Pt(10)


def clean_html_body(body):
    """Limpia HTML básico del cuerpo del correo."""
    if not body:
        return ""
    import re

    text = str(body)
    text = re.sub(r"<br\s*/?>", "\n", text, flags=re.IGNORECASE)
    text = re.sub(r"<p[^>]*>", "\n", text, flags=re.IGNORECASE)
    text = re.sub(r"</p>", "", text, flags=re.IGNORECASE)
    text = re.sub(r"<div[^>]*>", "\n", text, flags=re.IGNORECASE)
    text = re.sub(r"</div>", "", text, flags=re.IGNORECASE)
    text = re.sub(r"<[^>]+>", "", text)
    text = re.sub(r"&nbsp;", " ", text)
    text = re.sub(r"&amp;", "&", text)
    text = re.sub(r"&lt;", "<", text)
    text = re.sub(r"&gt;", ">", text)
    text = re.sub(r"&quot;", '"', text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def generate_docx(data, output_path):
    """Genera el documento Word a partir de los datos JSON."""
    entidad = data.get("entidad", "")
    periodo_inicio = data.get("periodo_inicio", "")
    periodo_fin = data.get("periodo_fin", "")
    cuenta = data.get("cuenta", "")
    total = data.get("total", 0)
    correos = data.get("correos", [])

    doc = Document()
    set_document_language(doc, "es-CO")

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)
    style.paragraph_format.space_after = Pt(4)

    sections = doc.sections
    for section in sections:
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)

    title = doc.add_heading(f"CORREOS ENVIADOS - {entidad}", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(
        f"Período: {periodo_inicio} a {periodo_fin}  |  Cuenta: {cuenta}  |  Total: {total} correos"
    )
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    add_horizontal_line(doc)

    for correo in correos:
        num = correo.get("num", "")
        asunto = correo.get("asunto", "(Sin asunto)")
        de = correo.get("de", cuenta)
        para = correo.get("para", "")
        cc = correo.get("cc", "")
        fecha = correo.get("fecha", "")
        cuerpo = clean_html_body(correo.get("cuerpo", ""))
        hilo = correo.get("hilo", [])

        heading = doc.add_heading(f"{num}. {asunto}", level=2)
        heading.paragraph_format.space_before = Pt(12)

        add_metadata_line(doc, "De", de)
        add_metadata_line(doc, "Para", para)
        if cc:
            add_metadata_line(doc, "CC", cc)
        add_metadata_line(doc, "Fecha", fecha)

        p_sep = doc.add_paragraph()
        p_sep.paragraph_format.space_before = Pt(6)

        add_body_text(doc, cuerpo)

        if hilo:
            for msg in hilo:
                add_horizontal_line(doc)
                msg_de = msg.get("de", "")
                msg_para = msg.get("para", "")
                msg_fecha = msg.get("fecha", "")
                msg_cuerpo = clean_html_body(msg.get("cuerpo", ""))

                p_reply_header = doc.add_paragraph()
                p_reply_header.paragraph_format.left_indent = Inches(0.3)
                p_reply_header.paragraph_format.space_before = Pt(4)
                run_header = p_reply_header.add_run(
                    f"--- Mensaje anterior ({msg_fecha}) ---"
                )
                run_header.font.size = Pt(9)
                run_header.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
                run_header.italic = True

                add_metadata_line(doc, "De", msg_de, indent=True)
                add_metadata_line(doc, "Para", msg_para, indent=True)

                add_body_text(doc, msg_cuerpo, indent=True)

        add_horizontal_line(doc)

    os.makedirs(os.path.dirname(output_path) or ".", exist_ok=True)
    doc.save(output_path)
    print(
        f"DOCX generado: {output_path} ({len(correos)} correos)", file=sys.stderr
    )


def main():
    if len(sys.argv) < 3:
        print("Uso: python3 scripts/correos_to_docx.py <json_input> <docx_output>")
        sys.exit(1)

    json_input = sys.argv[1]
    docx_output = sys.argv[2]

    with open(json_input, "r", encoding="utf-8") as f:
        data = json.load(f)

    generate_docx(data, docx_output)


if __name__ == "__main__":
    main()
