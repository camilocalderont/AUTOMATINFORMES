#!/usr/bin/env python3
"""
Genera un documento Word con informe de commits a partir de datos JSON.

Uso:
    python3 scripts/commits_to_docx.py <json_input> <docx_output>

Input JSON:
    {
        "entidad": "IDARTES",
        "periodo_inicio": "2026-02-01",
        "periodo_fin": "2026-02-28",
        "total_commits": 15,
        "repositorios": ["pandora_proxy", "pandora_l11"],
        "descripcion": "Durante el periodo se realizaron actividades de desarrollo...",
        "semanas": [
            {
                "nombre": "Semana 1 (01/02 - 07/02)",
                "estado": "Finalizado",
                "fecha_inicio": "01/02/2026",
                "fecha_fin": "07/02/2026",
                "funcionalidades": ["Actualizacion Dockerfile", "Ajuste sync-cdp.sh"]
            }
        ],
        "commits": [
            {
                "hash": "c92f8fc9",
                "fecha": "2026-02-05",
                "mensaje": "Actualizacion Dockerfile (Samba, sync-cdp.sh)",
                "rama": "master_proxy",
                "repositorio": "pandora_proxy"
            }
        ]
    }
"""

import json
import os
import sys
from datetime import datetime

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


def set_document_language(doc, lang_code="es-CO"):
    """Configura el idioma del documento Word a español (Colombia)."""
    styles_element = doc.styles.element
    rpr_default = styles_element.find(qn("w:docDefaults"))
    if rpr_default is None:
        from docx.oxml import OxmlElement
        rpr_default = OxmlElement("w:docDefaults")
        styles_element.insert(0, rpr_default)
    rpr = rpr_default.find(qn("w:rPrDefault"))
    if rpr is None:
        from docx.oxml import OxmlElement
        rpr = OxmlElement("w:rPrDefault")
        rpr_default.append(rpr)
    rpr_inner = rpr.find(qn("w:rPr"))
    if rpr_inner is None:
        from docx.oxml import OxmlElement
        rpr_inner = OxmlElement("w:rPr")
        rpr.append(rpr_inner)
    lang = rpr_inner.find(qn("w:lang"))
    if lang is None:
        from docx.oxml import OxmlElement
        lang = OxmlElement("w:lang")
        rpr_inner.append(lang)
    lang.set(qn("w:val"), lang_code)
    lang.set(qn("w:eastAsia"), lang_code)
    lang.set(qn("w:bidi"), lang_code)


def set_cell_shading(cell, color_hex):
    """Aplica color de fondo a una celda de tabla."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    shading = OxmlElement("w:shd")
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:color"), "auto")
    shading.set(qn("w:fill"), color_hex)
    cell._tc.get_or_add_tcPr().append(shading)


def generate_docx(data, output_path):
    """Genera el documento Word a partir de los datos JSON."""
    entidad = data.get("entidad", "")
    periodo_inicio = data.get("periodo_inicio", "")
    periodo_fin = data.get("periodo_fin", "")
    total_commits = data.get("total_commits", 0)
    repositorios = data.get("repositorios", [])
    descripcion = data.get("descripcion", "")
    semanas = data.get("semanas", [])
    commits = data.get("commits", [])

    doc = Document()
    set_document_language(doc, "es-CO")

    # Estilos del documento
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)

    # Título
    title = doc.add_heading(level=0)
    run = title.add_run(f"INFORME DE COMMITS - {entidad}")
    run.font.color.rgb = RGBColor(0, 51, 102)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Subtítulo con período y estadísticas
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(f"Período: {periodo_inicio} a {periodo_fin}")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(100, 100, 100)

    stats_para = doc.add_paragraph()
    stats_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    stats_para.add_run(f"Total de commits: {total_commits}").bold = True
    if repositorios:
        stats_para.add_run(f"  |  Repositorios: {', '.join(repositorios)}")

    doc.add_paragraph()

    # Descripción general
    if descripcion:
        doc.add_heading("Descripción General", level=1)
        for parrafo in descripcion.split("\n\n"):
            parrafo = parrafo.strip()
            if parrafo:
                doc.add_paragraph(parrafo)

    # Desarrollos semana a semana
    if semanas:
        doc.add_heading("Desarrollos Semana a Semana", level=1)

        for semana in semanas:
            nombre = semana.get("nombre", "")
            estado = semana.get("estado", "Finalizado")
            fecha_inicio = semana.get("fecha_inicio", "")
            fecha_fin = semana.get("fecha_fin", "")
            funcionalidades = semana.get("funcionalidades", [])

            doc.add_heading(nombre, level=2)

            # Metadata de la semana
            meta = doc.add_paragraph()
            meta.add_run("Estado: ").bold = True
            meta.add_run(estado)
            meta.add_run("\n")
            meta.add_run("Fecha Inicio: ").bold = True
            meta.add_run(fecha_inicio)
            meta.add_run("\n")
            meta.add_run("Fecha Fin: ").bold = True
            meta.add_run(fecha_fin)

            # Funcionalidades
            if funcionalidades:
                func_heading = doc.add_paragraph()
                func_heading.add_run("Funcionalidades:").bold = True
                for func in funcionalidades:
                    doc.add_paragraph(func, style="List Bullet")

    # Tabla completa de commits
    if commits:
        doc.add_heading("Tabla Completa de Commits", level=1)

        headers = ["#", "Commit ID", "Repositorio", "Rama", "Descripción", "Fecha"]
        table = doc.add_table(rows=1, cols=len(headers))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = "Table Grid"

        # Header
        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = header
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(255, 255, 255)
                    run.font.size = Pt(9)
            set_cell_shading(cell, "003366")

        # Filas con sombreado alternado
        for idx, commit in enumerate(commits, 1):
            row = table.add_row()
            values = [
                str(idx),
                str(commit.get("hash", "")),
                str(commit.get("repositorio", "")),
                str(commit.get("rama", "")),
                str(commit.get("mensaje", "")),
                str(commit.get("fecha", "")),
            ]
            for i, val in enumerate(values):
                cell = row.cells[i]
                cell.text = val
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(9)
            # Sombreado alternado
            if idx % 2 == 0:
                for i in range(len(headers)):
                    set_cell_shading(row.cells[i], "F5F5F5")

        # Ajustar anchos
        widths = [Cm(1), Cm(2.5), Cm(3), Cm(3.5), Cm(6), Cm(2.5)]
        for row in table.rows:
            for i, width in enumerate(widths):
                row.cells[i].width = width

    # Footer con timestamp
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run(
        f"Generado automáticamente el {datetime.now().strftime('%Y-%m-%d %H:%M')}"
    )
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(150, 150, 150)
    run.italic = True

    os.makedirs(os.path.dirname(output_path) or ".", exist_ok=True)
    doc.save(output_path)
    print(f"DOCX generado: {output_path} ({total_commits} commits)", file=sys.stderr)


def main():
    if len(sys.argv) < 3:
        print("Uso: python3 scripts/commits_to_docx.py <json_input> <docx_output>")
        sys.exit(1)

    json_input = sys.argv[1]
    docx_output = sys.argv[2]

    with open(json_input, "r", encoding="utf-8") as f:
        data = json.load(f)

    generate_docx(data, docx_output)


if __name__ == "__main__":
    main()
